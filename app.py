import os
import io
import re
import time
from dataclasses import dataclass
from typing import List, Dict, Any

import streamlit as st
import yaml
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Inches

# ---- OpenAI client (new SDK) ----
try:
    from openai import OpenAI
except Exception:
    OpenAI = None

# ------------------ Utilities ------------------
PHASES_DEFAULT = [
    "Clarify",
    "Probe Assumptions",
    "Evidence & Methods",
    "Counterexamples",
    "Transfer to Practice",
]

SYSTEM_PROMPT_DEFAULT = (
    "You are a Socratic dialogue coach for graduate students analyzing a specific assigned reading.\n"
    "Non-negotiables:\n"
    "- Always begin by asking the student to name the article (title/author) and a goal for the session.\n"
    "- Use one short, targeted question at a time.\n"
    "- Move in rounds: Clarify → Probe Assumptions → Evidence & Methods → Counterexamples → Transfer to Practice.\n"
    "- Anchor to page numbers when the student provides them; ask students to quote/paraphrase.\n"
    "- When asked to summarize, guide the student to build it: thesis, claims, evidence, implications.\n"
    "- If stuck, offer two or three next-step choices (as questions).\n"
    "- Avoid evaluative language about quality; focus on reasoning moves (definitions, warrants, evidence, limits, generalizability).\n"
    "- Surface equity/justice implications: ask how claims affect historically marginalized groups.\n"
    "- End with a 4-part wrap-up the student writes with your guidance: (1) thesis, (2) 2–3 claims+evidence, (3) 1 limitation/counterpoint, (4) 1 transfer to their practice.\n"
)

# ------------------ Config ------------------
@dataclass
class AppConfig:
    model: str
    phases: List[str]
    require_page_citations: bool
    max_turns: int
    equity_prompt: str


def load_config() -> AppConfig:
    # Load config.yaml if present; otherwise defaults
    cfg = {
        "model": "gpt-4o-mini",
        "phases": PHASES_DEFAULT,
        "require_page_citations": True,
        "max_turns": 30,
        "equity_prompt": (
            "As we go, consider equity: Who benefits or is burdened by the authors' claims? "
            "How might recommendations be reframed to center inclusion/belonging (e.g., multilingual learners, students with disabilities, racially/gender-marginalized groups)?"
        ),
    }
    if os.path.exists("config.yaml"):
        with open("config.yaml", "r") as f:
            try:
                y = yaml.safe_load(f)
                cfg.update({k: v for k, v in (y or {}).items() if v is not None})
            except Exception:
                pass
    return AppConfig(**cfg)


# ------------------ PDF Handling ------------------
@dataclass
class PDFContext:
    text_by_page: Dict[int, str]
    metadata: Dict[str, Any]


def extract_pdf_text(file_bytes: bytes) -> PDFContext:
    reader = PdfReader(io.BytesIO(file_bytes))
    text_by_page = {}
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        text_by_page[i] = text
    meta = reader.metadata or {}
    return PDFContext(text_by_page=text_by_page, metadata=meta)


def get_page_excerpt(ctx: PDFContext, page_num: int, n_chars: int = 700) -> str:
    text = ctx.text_by_page.get(page_num, "").strip()
    if not text:
        return "(No text extracted for that page; if your PDF is a scan, ensure it's OCR'd.)"
    # Return the first ~n_chars with ellipsis
    snippet = text[:n_chars].replace("\n", " ")
    if len(text) > n_chars:
        snippet += " …"
    return snippet


# ------------------ OpenAI Helpers ------------------

def get_client():
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        st.error("OPENAI_API_KEY not set in environment or Streamlit secrets.")
        st.stop()
    if OpenAI is None:
        st.error("openai SDK not installed. Add 'openai>=1.40.0' to requirements.txt")
        st.stop()
    return OpenAI()


def model_chat(client, model: str, messages: List[Dict[str, str]]) -> str:
    # Uses the Responses API if available; falls back to Chat Completions style
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0.4,
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        return f"[Model error] {e}"


# ------------------ Socratic Flow ------------------

def init_state(cfg: AppConfig):
    if "cfg" not in st.session_state:
        st.session_state.cfg = cfg
    st.session_state.setdefault("pdf_ctx", None)
    st.session_state.setdefault("article_title", "")
    st.session_state.setdefault("article_author", "")
    st.session_state.setdefault("session_goal", "")
    st.session_state.setdefault("doi", "")
    st.session_state.setdefault("phase_idx", 0)
    st.session_state.setdefault("turns", [])  # list of dicts: {role, content}
    st.session_state.setdefault("wrapup", {"thesis":"", "claims":"", "limitation":"", "transfer":""})
    st.session_state.setdefault("awaiting_answer", False)
    st.session_state.setdefault("last_question", "")


def reset_session():
    for k in ["pdf_ctx", "article_title", "article_author", "session_goal", "doi", "phase_idx", "turns", "wrapup", "awaiting_answer", "last_question"]:
        st.session_state.pop(k, None)
    init_state(st.session_state.cfg)


def require_page_citation(text: str) -> bool:
    # Accept patterns like p. 12, pp. 12-13, (p. 7), (pp. 5–6), page 3
    pattern = r"(p\.?\s*\d+|pp\.?\s*\d+\s*[–-]\s*\d+|page\s*\d+)"
    return re.search(pattern, text, flags=re.IGNORECASE) is not None


def next_phase_label() -> str:
    phases = st.session_state.cfg.phases
    idx = st.session_state.phase_idx
    return phases[idx % len(phases)]


def socratic_prompt_builder() -> List[Dict[str, str]]:
    cfg = st.session_state.cfg

    sys_prompt = SYSTEM_PROMPT_DEFAULT
    msgs = [{"role": "system", "content": sys_prompt}]

    # Context preamble (non-identifying)
    meta_lines = []
    if st.session_state.article_title:
        meta_lines.append(f"Article: {st.session_state.article_title}")
    if st.session_state.article_author:
        meta_lines.append(f"Author(s): {st.session_state.article_author}")
    if st.session_state.doi:
        meta_lines.append(f"DOI: {st.session_state.doi}")
    if st.session_state.session_goal:
        meta_lines.append(f"Student goal: {st.session_state.session_goal}")
    meta_lines.append(f"Current phase: {next_phase_label()}")
    meta_lines.append(cfg.equity_prompt)

    msgs.append({"role": "system", "content": "\n".join(meta_lines)})

    # Include transcript so far (kept in memory only; not stored server-side)
    for t in st.session_state.turns[-12:]:  # limit history tokens
        msgs.append({"role": t["role"], "content": t["content"]})

    # Instruction to ask exactly one question
    msgs.append({
        "role": "system",
        "content": (
            "Ask exactly ONE concise Socratic question next. Do not answer it. "
            "Prefer page-anchored prompts if the student has cited pages."
        )
    })
    return msgs


def ask_next_question(client):
    cfg = st.session_state.cfg
    messages = socratic_prompt_builder()
    model_reply = model_chat(client, cfg.model, messages)
    st.session_state.turns.append({"role": "assistant", "content": model_reply})
    st.session_state.last_question = model_reply
    st.session_state.awaiting_answer = True


def submit_student_answer(answer: str):
    st.session_state.turns.append({"role": "user", "content": answer})
    st.session_state.awaiting_answer = False
    # Advance phase every two Q&A turns to keep momentum
    st.session_state.phase_idx = (st.session_state.phase_idx + 1) % len(st.session_state.cfg.phases)


def build_wrapup(client):
    cfg = st.session_state.cfg
    transcript = "\n".join([f"{t['role'].upper()}: {t['content']}" for t in st.session_state.turns])
    prompt = (
        "Based on the conversation transcript, guide the student to draft a 4-part wrap-up: "
        "(1) Thesis in one sentence. (2) 2–3 key claims with page-anchored evidence. "
        "(3) 1 limitation or counterpoint. (4) 1 transfer to their practice. "
        "Respond with labeled sections Thesis, Claims, Limitation, Transfer and keep it concise."
    )
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT_DEFAULT},
        {"role": "user", "content": transcript + "\n\n" + prompt},
    ]
    draft = model_chat(client, cfg.model, messages)
    # Try to parse sections
    w = st.session_state.wrapup
    for key in ["thesis", "claims", "limitation", "transfer"]:
        w[key] = w.get(key, "")
    # Quick heuristic splits
    def extract(section_name):
        m = re.search(section_name + r"\s*[:\n]\s*(.*)", draft, re.IGNORECASE | re.DOTALL)
        return (m.group(1).strip() if m else "").split("\n\n")[0]
    st.session_state.wrapup["thesis"] = extract("Thesis") or st.session_state.wrapup["thesis"]
    st.session_state.wrapup["claims"] = extract("Claims") or st.session_state.wrapup["claims"]
    st.session_state.wrapup["limitation"] = extract("Limitation") or st.session_state.wrapup["limitation"]
    st.session_state.wrapup["transfer"] = extract("Transfer") or st.session_state.wrapup["transfer"]


def export_docx() -> bytes:
    w = st.session_state.wrapup
    doc = Document()
    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    doc.add_heading('Reading Wrap-Up', level=1)
    meta = []
    if st.session_state.article_title:
        meta.append(f"Article: {st.session_state.article_title}")
    if st.session_state.article_author:
        meta.append(f"Author(s): {st.session_state.article_author}")
    if st.session_state.doi:
        meta.append(f"DOI: {st.session_state.doi}")
    if st.session_state.session_goal:
        meta.append(f"Goal: {st.session_state.session_goal}")
    if meta:
        p = doc.add_paragraph("\n".join(meta))

    doc.add_heading('1) Thesis', level=2)
    doc.add_paragraph(w.get("thesis", ""))

    doc.add_heading('2) Claims + Evidence', level=2)
